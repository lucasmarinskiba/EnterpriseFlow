import streamlit as st
import pandas as pd
import sqlite3
import hashlib
import uuid
import numpy as np
import datetime
import os
import stripe
from fpdf import FPDF  # Importación corregida aquí
from database import DatabaseManager
from pathlib import Path
from payment_handler import PaymentHandler
from tensorflow.keras.models import load_model
import spacy
import smtplib
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from email.mime.base import MIMEBase
from email import encoders
import time
import requests

# Manejo de dependencias opcionales
try:
    import spacy
    from database import DatabaseManager
    from payment_handler import PaymentHandler
    from tensorflow.keras.models import load_model
except ImportError as e:
    st.error(f"Error de dependencias: {str(e)}")
    st.stop()

# Configuración inicial
st.set_page_config(
    page_title="EnterpriseFlow",
    page_icon="🏢",
    layout="wide"
)

class EnterpriseFlowApp:
    def __init__(self):
        try:
            self.nlp = spacy.load("es_core_news_sm")
        except Exception as e:
            st.error(f"Error cargando modelos de NLP: {str(e)}")
            st.info("Ejecuta: python -m spacy download es_core_news_sm")
            st.stop()
        
        self.db = DatabaseManager()
        self.payment = PaymentHandler()
        
        if 'logged_in' not in st.session_state:
            st.session_state.logged_in = False
        if 'current_user' not in st.session_state:
            st.session_state.current_user = None
            
        self._setup_ui()

    def _setup_ui(self):
        st.sidebar.image("https://via.placeholder.com/200x50.png?text=EnterpriseFlow", width=200)
        if not st.session_state.logged_in:
            self._show_login()
        else:
            self._rewards_header()
            self._show_main_interface()

    def _rewards_header(self):
        user = st.session_state.current_user

        # Ejemplo: estos datos deberían venir de tu base de datos o tu lógica real de avance
        # Simulación para el ejemplo (reemplaza los valores aquí por tus consultas reales):
        tareas_completadas = self.db.get_completed_tasks_count(user) if hasattr(self.db, "get_completed_tasks_count") else 12
        logros = self.db.get_user_achievements(user) if hasattr(self.db, "get_user_achievements") else ["Inicio de sesión diario", "Primer tarea completada"]
        dias_constancia = self.db.get_user_streak_days(user) if hasattr(self.db, "get_user_streak_days") else 7

        # Lógica de puntos y recompensas automática:
        puntos = tareas_completadas * 100 + len(logros) * 250 + dias_constancia * 10
        nivel = max(1, puntos // 500 + 1)
        # Insignias: 1 por cada 5 tareas completadas + 1 por cada 1000 puntos + 1 por cada 7 días de constancia
        insignias = (tareas_completadas // 5) + (puntos // 1000) + (dias_constancia // 7)

        st.markdown(
            f"""
            <div style="
                background-color:#f7f7fa;
                border-radius:8px;
                padding:0.2rem 1rem 0.1rem 1rem;
                margin-bottom:10px;
                margin-top:0px;
                box-shadow:0 1px 2px rgba(0,0,0,0.04);
                border:1px solid #ededed;
                ">
                <div style="display:flex;align-items:center;gap:10px;">
                    <span style="font-size:1.35rem;margin-right:6px;">🎮</span>
                    <b style="font-size:1.04rem;">Sistema de Recompensas</b>
                    <span style="flex:1;"></span>
                    <span style="font-size:0.97rem;display:flex;align-items:center;gap:2px;">
                        🥇
                        <span style="margin-left:2px;margin-right:10px;"><b>{puntos:,}</b></span>
                        🌟
                        <span style="margin-left:2px;margin-right:10px;"><b>{nivel}</b></span>
                        🏆
                        <span style="margin-left:2px;"><b>{insignias}</b></span>
                    </span>
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )

    def update_user_rewards(user_email, puntos_delta=0, tareas_completadas=0, dias_constancia=0, insignias_delta=0):
        conn = sqlite3.connect("enterprise_flow.db")
        c = conn.cursor()
        c.execute("SELECT puntos, nivel, insignias, tareas_completadas, dias_constancia FROM user_rewards WHERE user_email=?", (user_email,))
        row = c.fetchone()
        if row:
            puntos, nivel, insignias, tareas, dias = row
            puntos += puntos_delta
            tareas += tareas_completadas
            dias += dias_constancia
            nivel = max(1, puntos // 500 + 1)
            insignias += insignias_delta
            c.execute("""
                UPDATE user_rewards
                SET puntos=?, nivel=?, insignias=?, tareas_completadas=?, dias_constancia=?, updated_at=CURRENT_TIMESTAMP
                WHERE user_email=?
            """, (puntos, nivel, insignias, tareas, dias, user_email))
        else:
            c.execute("""
                INSERT INTO user_rewards (user_email, puntos, nivel, insignias, tareas_completadas, dias_constancia)
                VALUES (?, ?, ?, ?, ?, ?)
            """, (user_email, puntos_delta, 1, insignias_delta, tareas_completadas, dias_constancia))
        conn.commit()
        conn.close()
    
    def _show_login(self):
        with st.sidebar:
            st.header("Bienvenido a EnterpriseFlow")
            tab1, tab2 = st.tabs(["Iniciar Sesión", "Registrarse"])
            
            with tab1:
                email_login = st.text_input("Correo electrónico")
                password_login = st.text_input("Contraseña", type="password")
                if st.button("Ingresar"):
                    if self.db.verify_user(email_login, password_login):
                        st.session_state.logged_in = True
                        st.session_state.current_user = email_login
                        st.rerun()
                    else:
                        st.error("Credenciales incorrectas")
            
            with tab2:
                email_register = st.text_input("Correo para registro")
                password_register = st.text_input("Contraseña nueva", type="password")
                if st.button("Crear Cuenta"):
                    try:
                        self.db.create_user(email_register, password_register)
                        st.success("¡Cuenta creada exitosamente!")
                    except sqlite3.IntegrityError:
                        st.error("Este correo ya está registrado")

    def _show_main_interface(self):
        menu = st.sidebar.radio(
            "Menú Principal",
            ["🏠 Inicio", "📁 Documentos", "🤖 Automatización", "😌 Bienestar", "🎓 Aprendizaje", "🔒 Feedback Anónimo", "⚖️ Cumplimiento", "💳 Suscripción"]
        )
        
        if menu == "🏠 Inicio":
            self._show_dashboard()
        elif menu == "📁 Documentos":
            self._show_documents()
        elif menu == "🤖 Automatización":
            self._show_automation()
        elif menu == "😌 Bienestar":
            self._show_wellness()
        elif menu == "🎓 Aprendizaje":
            self._learning_portal()
        elif menu == "🔒 Feedback Anónimo":
            self._show_feedback_system()
        elif menu == "⚖️ Cumplimiento":
            self._show_compliance()
        elif menu == "💳 Suscripción":
            self._show_payment()

    # app/sections/automation.py
    def show_automation_preview(config):
        with st.expander("Vista Previa del Flujo"):
            st.graphviz_chart(f"""
            digraph {{
               {config['trigger']} -> {config['action']}
               {config['action']} -> {config['notification']}
            }}
            """)
            
    def _show_dashboard(self):
        st.title("Panel de Control")
        st.write(f"Bienvenido: [{st.session_state.current_user}](mailto:{st.session_state.current_user})")

        # --- SECCIÓN TUTORIALES ---
        st.markdown("---")
        st.header("📖 Tutoriales: Introducción a como usar EnterpriseFlow")
         
        with st.expander("🤖 Automatización de Tareas"):
            st.markdown("""
            1. Ve al menú lateral y selecciona **Automatización**.
            2. Completa los campos para crear una factura, programar tareas o configurar envíos de emails masivos.
            3. Haz clic en los botones correspondientes para guardar o ejecutar cada acción.
            4. Observa la vista previa de tus automatizaciones y revisa los resultados en tiempo real.
            """)

        with st.expander("😌 Bienestar del Equipo"):
            st.markdown("""
            1. Accede a la sección **Bienestar** en el menú.
            2. Usa el panel para calcular el riesgo de burnout y registra reconocimientos a tus colegas.
            3. Explora los módulos de salud, descansos programados, aprendizaje y metas personales.
            4. Personaliza y guarda tus propios datos de salud para llevar tu propio registro.
            """)

        with st.expander("🔒 Feedback Anónimo"):
            st.markdown("""
            1. Selecciona **Feedback Anónimo** en el menú lateral.
            2. Elige el tipo de feedback y escribe tu mensaje de forma confidencial.
            3. Haz clic en **Enviar Feedback** para que la administración lo reciba sin conocer tu identidad.
            """)

        with st.expander("⚖️ Cumplimiento y Auditoría"):
            st.markdown("""
            1. En el menú, haz clic en **Cumplimiento**.
            2. Sube documentos para auditoría normativa (GDPR, SOX, ISO27001, etc.).
            3. Revisa los resultados automáticos y asegúrate de que tus archivos cumplen con los estándares requeridos.
            """)

        st.info("¿Necesitas más ayuda? Contacta a soporte o revisa la documentación oficial para ver videos y guías paso a paso.")

    def _show_documents(self):
        st.header("📁 Gestor de Documentos")
        st.markdown("Sube, escanea, entrega y gestiona tus documentos digitales.")

        # Subida de documentos
        uploaded_file = st.file_uploader("Sube un documento (PDF, imagen, Word)", type=["pdf", "png", "jpg", "jpeg", "docx"])
        user = st.session_state.current_user

        if uploaded_file:
            # Guardar el archivo temporalmente o en base de datos
            save_path = f"uploaded_docs/{user}_{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}_{uploaded_file.name}"
            os.makedirs(os.path.dirname(save_path), exist_ok=True)
            with open(save_path, "wb") as f:
                f.write(uploaded_file.getbuffer())
            st.success(f"Documento '{uploaded_file.name}' guardado correctamente.")

            if uploaded_file.type == "application/pdf":
                with st.expander("📄 Escanear PDF (extraer texto)"):
                    import PyPDF2
                    reader = PyPDF2.PdfReader(save_path)
                    text = "\n".join([page.extract_text() or "" for page in reader.pages])
                    st.text_area("Texto extraído", value=text, height=200)
            elif "image" in uploaded_file.type:
                with st.expander("🔎 Escanear Imagen (OCR)"):
                    try:
                        import pytesseract
                        from PIL import Image
                        img = Image.open(save_path)
                        st.image(img, caption="Imagen subida", use_column_width=True)
                        text = pytesseract.image_to_string(img, lang="spa")
                        st.text_area("Texto extraído (OCR)", value=text, height=200)
                    except ImportError:
                        st.warning("pytesseract y pillow necesarios para OCR de imagen. Instálalos con pip si quieres esta función.")
            elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                with st.expander("📄 Leer Word"):
                    try:
                        from docx import Document
                        doc = Document(save_path)
                        text = "\n".join([para.text for para in doc.paragraphs])
                        st.text_area("Texto extraído", value=text, height=200)
                    except ImportError:
                        st.warning("python-docx necesario para abrir archivos Word.")

        # Listar documentos subidos por el usuario
        st.markdown("### Tus documentos subidos")
        doc_folder = "uploaded_docs"
        docs = []
        if os.path.exists(doc_folder):
            docs = [f for f in os.listdir(doc_folder) if f.startswith(user+"_")]
        if docs:
            for docfile in docs:
                st.write(f"📄 {docfile}")
                with open(f"{doc_folder}/{docfile}", "rb") as f:
                    st.download_button(
                        label="Descargar",
                        data=f.read(),
                        file_name=docfile.split("_",2)[-1],
                        mime="application/octet-stream"
                    )
        else:
            st.info("Aún no has subido documentos.")

        st.markdown("---")
        st.header("🧾 Generar y Enviar Recibo Digital")

        with st.form("generate_receipt"):
            nombre = st.text_input("Nombre del receptor")
            monto = st.number_input("Monto", min_value=0.0)
            concepto = st.text_input("Concepto del recibo")
            email_receptor = st.text_input("Email para enviar recibo")
            submit = st.form_submit_button("Generar Recibo PDF y Enviar por Email")

            if submit and nombre and monto and concepto and email_receptor:
                # Crear PDF de recibo en memoria (¡NO en disco!)
                from fpdf import FPDF
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=14)
                pdf.cell(0, 10, "Recibo Digital", ln=1, align="C")
                pdf.set_font("Arial", size=12)
                pdf.ln(8)
                pdf.cell(0, 10, f"Recibí de: {nombre}", ln=1)
                pdf.cell(0, 10, f"Monto: ${monto:.2f}", ln=1)
                pdf.cell(0, 10, f"Concepto: {concepto}", ln=1)
                pdf.cell(0, 10, f"Emitido por: {user}", ln=1)
                pdf.cell(0, 10, f"Fecha: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}", ln=1)
                # Guardar como bytes en memoria — robusto para cualquier versión de FPDF
                pdf_output = pdf.output(dest='S')
                if isinstance(pdf_output, str):
                    pdf_bytes = pdf_output.encode('latin1')
                else:
                    pdf_bytes = pdf_output
                st.success("Recibo PDF generado correctamente.")
                st.download_button("Descargar recibo PDF", data=pdf_bytes, file_name="recibo.pdf", mime="application/pdf")

                # Enviar por email
                try:
                    import smtplib
                    from email.mime.multipart import MIMEMultipart
                    from email.mime.base import MIMEBase
                    from email.mime.text import MIMEText
                    from email import encoders

                    msg = MIMEMultipart()
                    msg["From"] = st.secrets["smtp"]["user"]
                    msg["To"] = email_receptor
                    msg["Subject"] = "Recibo Digital"

                    body = f"Estimado/a {nombre},\nAdjunto encontrará su recibo digital por el concepto: {concepto}."
                    msg.attach(MIMEText(body, "plain"))

                    part = MIMEBase("application", "octet-stream")
                    part.set_payload(pdf_bytes)
                    encoders.encode_base64(part)
                    part.add_header("Content-Disposition", f"attachment; filename=recibo.pdf")
                    msg.attach(part)

                    smtp_server = st.secrets["smtp"]["server"]
                    smtp_port = st.secrets["smtp"]["port"]
                    sender_email = st.secrets["smtp"]["user"]
                    sender_password = st.secrets["smtp"]["password"]

                    with smtplib.SMTP(smtp_server, smtp_port) as server:
                        server.starttls()
                        server.login(sender_email, sender_password)
                        server.sendmail(sender_email, email_receptor, msg.as_string())
                    st.success(f"Recibo enviado a {email_receptor} correctamente.")
                except Exception as e:
                    st.error(f"Error enviando email: {str(e)}")
    
    def _show_automation(self):
        with st.expander("🤖 Automatización de Tareas", expanded=True):
            col1, col2, col3 = st.columns(3)
            
            with col1:
                st.subheader("Generador de Facturas")
                client_name = st.text_input("Nombre del Cliente")
                subtotal = st.number_input("Subtotal", min_value=0.0)
                client_address = st.text_input("Dirección del Cliente")
                client_email = st.text_input("Correo Electrónico del Cliente")
                
                if st.button("Generar Factura"):
                    invoice_data = {
                        'client_name': client_name,
                        'subtotal': subtotal,
                        'client_address': client_address,
                        'client_email': client_email
                    }
                    
                    invoice = self._generate_invoice(invoice_data)
                    
                    smtp_server = st.secrets["smtp"]["server"]
                    smtp_port = st.secrets["smtp"]["port"]
                    sender_email = st.secrets["smtp"]["user"]
                    sender_password = st.secrets["smtp"]["password"]
                    
                    msg = MIMEMultipart()
                    msg['From'] = sender_email
                    msg['To'] = client_email
                    msg['Subject'] = f"Factura {invoice['invoice_number']} - {client_name}"
                    
                    body = f"""
                    Hola {client_name},
                    
                    Adjunto encontrará su factura número {invoice['invoice_number']}.
                    
                    Detalles:
                    - Total: ${invoice['total']}
                    - Fecha: {invoice['date']}
                    
                    Gracias por su preferencia.
                    """
                    msg.attach(MIMEText(body, 'plain'))
                    
                    filename = f"Factura_{invoice['invoice_number']}.pdf"
                    part = MIMEBase('application', 'octet-stream')
                    part.set_payload(invoice['pdf_data'])
                    encoders.encode_base64(part)
                    part.add_header('Content-Disposition', f'attachment; filename= {filename}')
                    msg.attach(part)
                    
                    try:
                        with smtplib.SMTP(smtp_server, smtp_port) as server:
                            server.starttls()
                            server.login(sender_email, sender_password)
                            server.sendmail(sender_email, client_email, msg.as_string())
                        st.success(f"Factura enviada a {client_email}!")
                    except Exception as e:
                        st.error(f"Error enviando email: {str(e)}")

            with col2:
                st.subheader("Programación de Tareas")
                task_type = st.selectbox("Tipo de Tarea", ["Reporte", "Recordatorio", "Backup"])
                schedule_time = st.time_input("Hora de Ejecución")
                
                if st.button("Programar Tarea"):
                    self.db.save_automation_task(st.session_state.current_user, {
                        'type': task_type,
                        'schedule': schedule_time.strftime("%H:%M")
                    })
                    st.success("Tarea programada exitosamente")

            with col3:
                st.subheader("Nuevas Automatizaciones")
                
                with st.container(border=True):
                    st.markdown("**📧 Email Masivo**")
                    email_subject = st.text_input("Asunto del Email")
                    email_template = st.text_area("Plantilla HTML")
                    if st.button("Programar Envío"):
                        self.db.save_automation_task(st.session_state.current_user, {
                            'type': 'email_masivo',
                            'subject': email_subject,
                            'template': email_template
                        })
                        st.success("Envío programado!")
                
                with st.container(border=True):
                    st.markdown("**🔄 Sync CRM**")
                    crm_action = st.selectbox("Acción", ["Actualizar clientes", "Importar leads"])
                    sync_frequency = st.selectbox("Frecuencia", ["Diario", "Semanal", "Mensual"])
                    if st.button("Configurar Sync"):
                        self.db.save_automation_task(st.session_state.current_user, {
                            'type': 'crm_sync',
                            'action': crm_action,
                            'frequency': sync_frequency
                        })
                        st.success("Sincronización configurada")

            with st.container():
                st.subheader("Automatizaciones Avanzadas")
                adv_col1, adv_col2 = st.columns(2)
                
                with adv_col1:
                    st.markdown("**🔮 Análisis Predictivo**")
                    model_type = st.selectbox("Modelo", ["Ventas", "Retención", "Inventario"])
                    if st.button("Ejecutar Modelo"):
                        self._run_predictive_model(model_type)
                        st.success("Modelo ejecutado")
                
                with adv_col2:
                    st.markdown("**⚙️ Integración Externa**")
                    api_endpoint = st.text_input("URL API")
                    if st.button("Conectar"):
                        self._test_api_connection(api_endpoint)
                        st.success("Conexión exitosa")

    def get_smtp_settings(email):
        domain = email.split('@')[-1].lower()
        if domain == "gmail.com":
            return "smtp.gmail.com", 587
        elif domain in ["hotmail.com", "outlook.com", "live.com"]:
            return "smtp.office365.com", 587
        elif domain == "yahoo.com":
            return "smtp.mail.yahoo.com", 587
        elif domain in ["icloud.com", "me.com"]:
            return "smtp.mail.me.com", 587
        elif domain == "zoho.com":
            return "smtp.zoho.com", 587
        else:
            # Por defecto, podrías pedir al usuario que lo indique
            raise ValueError("Proveedor de correo no soportado automáticamente")

    # Uso:
    sender_email = "usuario@hotmail.com"
    smtp_server, smtp_port = get_smtp_settings(sender_email)
    
    def _generate_invoice(self, data):
        iva_rate = 0.16 if 'MEX' in data['client_address'] else 0.21
        return {
            'client': data['client_name'],
            'total': round(data['subtotal'] * (1 + iva_rate), 2),
            'compliance_check': self._check_compliance(data),
            'invoice_number': f"INV-{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}",
            'date': datetime.datetime.now().strftime("%d/%m/%Y"),
            'pdf_data': b''
        }

    def _check_compliance(self, data):
        return True

    def _show_wellness(self):
        st.markdown("---")
        st.subheader("🩺 Ficha Médica del Empleado")
        user = st.session_state.current_user
        # Mostrar/actualizar ficha médica
        ficha = self.db.get_medical_record(user)
        with st.form("ficha_medica"):
            patologia = st.text_input("Patología principal", value=ficha.get("patologia", "") if ficha else "")
            enfermedades = st.text_area("Otras enfermedades", value=ficha.get("enfermedades", "") if ficha else "")
            embarazo = st.checkbox("Embarazo", value=bool(ficha.get("embarazo", 0)) if ficha else False)
            observaciones = st.text_area("Observaciones", value=ficha.get("observaciones", "") if ficha else "")
            if st.form_submit_button("Guardar ficha médica"):
                self.db.save_medical_record(user, patologia, enfermedades, embarazo, observaciones)
                st.success("Ficha médica actualizada.")

        st.markdown("---")
        st.subheader("📋 Faltas y Permisos de Salud")
        with st.form("solicitar_permiso"):
            tipo = st.selectbox("Tipo de permiso", ["Vacaciones", "Enfermedad", "Otro"])
            fecha_inicio = st.date_input("Desde")
            fecha_fin = st.date_input("Hasta")
            motivo = st.text_input("Motivo")
            observaciones = st.text_area("Observaciones")
            if st.form_submit_button("Solicitar permiso"):
                self.db.save_leave_request(user, tipo, fecha_inicio, fecha_fin, motivo, observaciones)
                st.success("Permiso solicitado.")

        # NUEVO: Apellido y nombre del empleado
        apellido = st.text_input("Apellido del empleado")
        nombre = st.text_input("Nombre del empleado")

        # Obtener ficha médica previa (si existe)
        ficha = self.db.get_medical_record(user)

        # NUEVO: Adjuntar archivo a la ficha médica
        uploaded_file = st.file_uploader("Adjunta un documento médico (PDF, imagen, Word)", type=["pdf", "png", "jpg", "jpeg", "docx"])

        with st.form("ficha_medica"):
            patologia = st.text_input("Patología principal", value=ficha.get("patologia", "") if ficha else "")
            enfermedades = st.text_area("Otras enfermedades", value=ficha.get("enfermedades", "") if ficha else "")
            embarazo = st.checkbox("Embarazo", value=bool(ficha.get("embarazo", 0)) if ficha else False)
            observaciones = st.text_area("Observaciones", value=ficha.get("observaciones", "") if ficha else "")

            guardar = st.form_submit_button("Guardar ficha médica")

            if guardar:
                # Guardar archivo en carpeta correspondiente si se adjuntó
                file_path = None
                if uploaded_file and apellido and nombre:
                    empleado_folder = f"fichas_medicas/{apellido}_{nombre}"
                    os.makedirs(empleado_folder, exist_ok=True)
                    file_path = f"{empleado_folder}/{datetime.datetime.now().strftime('%Y%m%d%H%M%S')}_{uploaded_file.name}"
                    with open(file_path, "wb") as f:
                        f.write(uploaded_file.getbuffer())

                # Guardar en base de datos (debes agregar la columna file_path en la tabla si aún no existe)
                self.db.save_medical_record(
                    user, patologia, enfermedades, embarazo, observaciones,
                    apellido=apellido, nombre=nombre, file_path=file_path
                )
                st.success("Ficha médica actualizada y archivo guardado correctamente.")

        # Mostrar archivos/fichas guardadas de ese empleado
        if apellido and nombre:
            empleado_folder = f"fichas_medicas/{apellido}_{nombre}"
            if os.path.exists(empleado_folder):
                archivos = os.listdir(empleado_folder)
                if archivos:
                    st.markdown("#### Archivos médicos guardados:")
                    for archivo in archivos:
                        st.write(f"📄 {archivo}")
                        with open(os.path.join(empleado_folder, archivo), "rb") as f:
                            st.download_button(
                                label="Descargar",
                                data=f.read(),
                                file_name=archivo,
                                mime="application/octet-stream"
                            )
                else:
                    st.info("No hay archivos médicos para este empleado.")
        
        # Mostrar historial
        st.markdown("### Permisos solicitados")
        leaves = self.db.get_leave_requests(user)
        if leaves:
            for lv in leaves:
                st.markdown(f"- **{lv['tipo_permiso']}**: {lv['fecha_inicio']} a {lv['fecha_fin']} ({lv['estado']})<br> Motivo: {lv['motivo']}<br>Obs: {lv['observaciones']}", unsafe_allow_html=True)
        else:
            st.info("Aún no has registrado permisos.")
        
        # ... el resto de tu módulo de bienestar ...

        with st.expander("😌 Bienestar del Equipo", expanded=True):
            col1, col2 = st.columns(2)
            
            with col1:
                st.subheader("Predicción de Burnout")
                hours_worked = st.slider("Horas trabajadas esta semana", 0, 100, 40)
                if st.button("Calcular Riesgo"):
                    prediction = self._predict_burnout(np.array([[hours_worked, 0, 0, 0, 0]]))
                    st.metric("Riesgo de Burnout", f"{prediction}%")

            with col2:
                st.subheader("Sistema de Reconocimiento")
                colleague = st.text_input("Nombre del Colega")
                colleague_email = st.text_input("Email del Colega")
                recognition = st.text_area("Mensaje de Reconocimiento")
                signing_authority = st.selectbox("Firmante", ["CEO", "Gerente General"])
                
                if st.button("Enviar 🏆"):
                    certificate_data = self._generate_certificate(
                        colleague=colleague,
                        recognition=recognition,
                        signer=signing_authority
                    )
                    cert_id = self.db.save_recognition(
                        user = st.session_state.current_user,
                        colleague = colleague,
                        recognition = recognition
                    )
                    if certificate_data and 'cert_id' in certificate_data:
                        certificate_id = certificate_data['cert_id'],
                        signer=signing_authority,
                        pdf_data = certificate_data['pdf_bytes']
                    
                    if self._send_recognition_email(colleague_email, certificate_data):
                        st.success(f"Certificado enviado a {colleague_email}!")
                        st.download_button(
                            label = "Descargar Certificado",
                            data = certificate_data['pdf_bytes'],
                            file_name = f"Certificado_{cert_id}.pdf",
                            mime = "application/pdf"
                        )
                    else:
                        st.error("Error enviando el certificado")

            st.markdown("---")
            self._health_dashboard()
            self._smart_breaks()
            self._personal_goals()
            self._meditation_module()
            self._team_network()
            self._workload_monitor()
            self._gamification_system()

    
    def _generate_certificate(self, colleague, recognition, signer):
        try:
            signer_key = signer.lower().replace(" ", "_")
            signature_path = st.secrets["signatures"][signer_key]
            full_path = BASE_DIR / signature_path
            
            if not full_path.exists():
                raise FileNotFoundError(f"❌ Archivo de firma no encontrado: {full_path}")

            # Crear PDF
            pdf = FPDF()
            pdf.add_page()
            pdf.set_auto_page_break(auto=True, margin=15)

            # Configuración del documento
            pdf.set_font("Arial", 'B', 16)
            pdf.cell(0, 10, "Certificado de Reconocimiento", ln=1, align='C')
            pdf.ln(15)

            # Contenido del certificado
            pdf.set_font("Arial", '', 12)
            pdf.multi_cell(0, 10, f"Se reconoce oficialmente a {colleague} por:", align='C')
            pdf.ln(10)
            pdf.set_font("Arial", 'I', 12)
            pdf.multi_cell(0, 8, f'"{recognition}"')
            pdf.ln(20)

            # Insertar firma
            pdf.image(str(full_path), x=50, w=30, h=15, type='PNG')
            pdf.set_font("Arial", 'I', 10)
            pdf.cell(0, 10, f"Firmado por: {signer}", ln=1, align='R')
            pdf.ln(15)

            # Pie de página
            cert_id = str(uuid.uuid4())[:8].upper()
            pdf.set_font("Arial", '', 8)
            pdf.cell(0, 10, f"ID de Certificado: {cert_id}", 0, 1, 'C')
            pdf.cell(0, 10, datetime.datetime.now().strftime("Emitido el %d/%m/%Y a las %H:%M"), 0, 1, 'C')

            return {
                'pdf_bytes': pdf.output(dest='S').encode('latin1'),
                'cert_id': cert_id
            }
            
        except KeyError as e:
            st.error(f"❌ Firma no configurada para '{signer}'. Verifica secrets.toml")
            return None
        except Exception as e:
            st.error(f"🚨 Error al generar certificado: {str(e)}")
            return None
    
    def _send_recognition_email(self, recipient, certificate_data):
        try:
            msg = MIMEMultipart()
            msg['From'] = st.secrets["smtp"]["user"]
            msg['To'] = recipient
            msg['Subject'] = "🏆 Reconocimiento Oficial - Tu Certificado"
            
            body = f"""
            ¡Felicitaciones!
            
            Has recibido un reconocimiento oficial de la empresa.
            Adjunto encontrarás tu certificado digital con validez oficial.
            
            ID del Certificado: {certificate_data['cert_id']}
            """
            msg.attach(MIMEText(body, 'plain'))
            
            part = MIMEBase('application', 'octet-stream')
            part.set_payload(certificate_data['pdf_bytes'])
            encoders.encode_base64(part)
            part.add_header('Content-Disposition', f'attachment; filename= "Certificado_{certificate_data["cert_id"]}.pdf"')
            msg.attach(part)
            
            with smtplib.SMTP(st.secrets["smtp"]["server"], st.secrets["smtp"]["port"]) as server:
                server.starttls()
                server.login(st.secrets["smtp"]["user"], st.secrets["smtp"]["password"])
                server.sendmail(st.secrets["smtp"]["user"], recipient, msg.as_string())
            return True
        except Exception as e:
            st.error(f"Error: {str(e)}")
            return False

    def _predict_burnout(self, input_data):
        try:
            return min(100, int(input_data[0][0] * 1.5))
        except Exception as e:
            st.error(f"Error: {str(e)}")
            return 0

    def _show_feedback_system(self):
        with st.expander("🔒 Sistema de Feedback Anónimo", expanded=True):
            feedback_type = st.selectbox("Tipo de Feedback", ["Para el equipo", "Para liderazgo", "Sugerencias generales"])
            feedback = st.text_area("Escribe tu feedback (máx. 500 caracteres)", max_chars=500)
            if st.button("Enviar Feedback"):
                self.db.save_anonymous_feedback(feedback_type, feedback)
                st.success("¡Gracias por tu contribución! Tu feedback es anónimo.")

    def _health_dashboard(self):
        with st.container(border=True):
            st.subheader("📊 Panel de Salud Integral")

            # Cargar los datos del usuario
            user = st.session_state.current_user
            health_data = self.db.get_health_data(user)
            # Si no existen, usa valores por defecto
            if not health_data:
                health_data = {'dias': 0, 'sueno': 0.0, 'pasos': 0}

            # Inputs para que el usuario pueda cambiarlos
            dias = st.number_input("Días sin incidentes", min_value=0, value=health_data['dias'])
            sueno = st.number_input("Horas Sueño Promedio", min_value=0.0, step=0.1, value=health_data['sueno'])
            pasos = st.number_input("Pasos Diarios", min_value=0, value=health_data['pasos'])

            if st.button("Guardar mis registros de salud"):
                self.db.save_health_data(user, dias, sueno, pasos)
                st.success("¡Datos guardados!")

            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("📅 Días sin incidentes", dias)
            with col2:
                st.metric("💤 Horas Sueño Promedio", sueno)
            with col3:
                st.metric("🚶 Pasos Diarios", pasos)

    def _smart_breaks(self):
        with st.container(border=True):
            st.subheader("⏰ Programador de Descansos Inteligentes")
            break_frequency = st.slider("Intervalo entre descansos (minutos)", 30, 120, 50)
            break_duration = st.slider("Duración del descanso (minutos)", 5, 15, 7)
            if st.button("Activar Recordatorios"):
                self._schedule_breaks(break_frequency, break_duration)

    def _schedule_breaks(self, frequency, duration):
        st.session_state.break_config = {
            'frequency': frequency,
            'duration': duration
        }
        st.success(f"Descansos programados cada {frequency} minutos por {duration} minutos")

    def _personal_goals(self):
        st.header("🏆 Sistema de Metas Personales")
        user = st.session_state.current_user

        with st.form(key="add_goal_form"):
            new_goal = st.text_input("Establece tu objetivo personal esta semana")
            submitted = st.form_submit_button("Guardar Objetivo")
            if submitted and new_goal.strip():
                self.db.save_personal_goal(user, new_goal.strip())
                st.success("¡Objetivo guardado!")

        st.markdown("#### Tus objetivos de esta semana:")

        goals = self.db.get_personal_goals(user)
        if not goals:
            st.info("Aún no tienes objetivos registrados.")
        else:
            for goal_id, goal_text in goals:
                col1, col2, col3 = st.columns([6,2,2])
                if "edit_goal" not in st.session_state:
                    st.session_state["edit_goal"] = {}

                # Mostrar en modo edición o como texto normal
                if st.session_state["edit_goal"].get(goal_id, False):
                    with col1:
                        edited_text = st.text_input(f"Edita objetivo {goal_id}", value=goal_text, key=f"edit_input_{goal_id}")
                    with col2:
                        if st.button("Guardar", key=f"save_btn_{goal_id}"):
                            self.db.edit_personal_goal(goal_id, edited_text)
                            st.session_state["edit_goal"][goal_id] = False
                            st.experimental_rerun()
                    with col3:
                        if st.button("Cancelar", key=f"cancel_btn_{goal_id}"):
                            st.session_state["edit_goal"][goal_id] = False
                            st.experimental_rerun()
                else:
                    with col1:
                        st.write(goal_text)
                    with col2:
                        if st.button("Editar", key=f"edit_btn_{goal_id}"):
                            st.session_state["edit_goal"][goal_id] = True
                            st.experimental_rerun()
                    with col3:
                        if st.button("Eliminar", key=f"delete_btn_{goal_id}"):
                            self.db.delete_personal_goal(goal_id)
                            st.experimental_rerun()

    def add_personal_goal(user_email, goal_text):
        conn = sqlite3.connect("enterprise_flow.db")
        c = conn.cursor()
        c.execute("INSERT INTO personal_goals (user_email, goal_text) VALUES (?, ?)", (user_email, goal_text))
        conn.commit()
        conn.close()
    
    def _meditation_module(self):
        with st.container(border=True):
            st.subheader("🧘 Sesiones de Relajación")
            duration_str = st.radio("Duración:", ["5 min", "10 min", "15 min"], key="meditation_duration")
            minutes = int(duration_str.split()[0])
            seconds = minutes * 60

            # Ruta según la duración (cada archivo debe existir en assets/)
            if minutes == 5:
                audio_path = "assets/meditacion.mp3"
            elif minutes == 10:
                audio_path = "assets/meditacion2.mp3"
            elif minutes == 15:
                audio_path = "assets/meditacion3.mp3"
            else:
                st.error("Duración no soportada.")
                return

            # Leer el archivo MP3 solo una vez por cambio de audio
            if (
                "audio_bytes" not in st.session_state
                or st.session_state.get("audio_path") != audio_path
            ):
                if os.path.exists(audio_path):
                    with open(audio_path, "rb") as f:
                        st.session_state["audio_bytes"] = f.read()
                    st.session_state["audio_path"] = audio_path
                else:
                    st.session_state["audio_bytes"] = None
                    st.session_state["audio_path"] = None

            if st.session_state["audio_bytes"] is None:
                st.error("No se encontró el archivo de audio de meditación para esta duración.")
                return

            if st.button("Iniciar Meditación Guiada", key="meditation_start_btn"):
                st.session_state["meditation_active"] = True
                st.session_state["meditation_start_time"] = time.time()
                st.session_state["meditation_total_seconds"] = seconds

            if st.session_state.get("meditation_active", False):
                elapsed = int(time.time() - st.session_state["meditation_start_time"])
                remaining = st.session_state["meditation_total_seconds"] - elapsed
                mins, secs = divmod(max(remaining, 0), 60)
                st.markdown(f"### ⏳ Tiempo restante: {mins:02d}:{secs:02d}")

                st.audio(st.session_state["audio_bytes"], format="audio/mp3")
                st.download_button(
                    label="Descargar música de meditación (MP3)",
                    data=st.session_state["audio_bytes"],
                    file_name=os.path.basename(audio_path),
                    mime="audio/mp3"
                )

                if remaining > 0:
                    st.experimental_rerun()
                else:
                    st.session_state["meditation_active"] = False
                    st.success("¡La sesión ha finalizado! Puedes abrir los ojos y continuar tu día.")
    
    def _team_network(self):
        import json
        st.subheader("👥 Mapa de Relaciones del Equipo")

        user_key = f"team_graph_{st.session_state.current_user}"
        if user_key not in st.session_state:
            # Estructura inicial por defecto
            st.session_state[user_key] = {
                "nodes": ["CEO", "Gerente", "Equipo A", "Equipo B", "Miembro 1", "Miembro 2", "Miembro 3"],
                "edges": [
                    ("CEO", "Gerente"),
                    ("Gerente", "Equipo A"),
                    ("Gerente", "Equipo B"),
                    ("Equipo A", "Miembro 1"),
                    ("Equipo A", "Miembro 2"),
                    ("Equipo B", "Miembro 3"),
                ]
            }

        # Estado de edición
        edit_state_key = f"edit_team_graph_{st.session_state.current_user}"
        if edit_state_key not in st.session_state:
            st.session_state[edit_state_key] = False

        # Botón para alternar edición
        if st.button("Editar" if not st.session_state[edit_state_key] else "Terminar Edición", key="edit_team_graph_btn"):
             st.session_state[edit_state_key] = not st.session_state[edit_state_key]

        graph = st.session_state[user_key]

        if st.session_state[edit_state_key]:
            st.markdown("#### Nodos (Personas/Equipos)")
            with st.form("add_node_form"):
                new_node = st.text_input("Agregar nuevo nodo (nombre)")
                add_node = st.form_submit_button("Agregar Nodo")
                if add_node and new_node.strip():
                    if new_node not in graph["nodes"]:
                        graph["nodes"].append(new_node.strip())
                    else:
                        st.warning("Ese nodo ya existe.")

            if len(graph["nodes"]) > 0:
                node_to_remove = st.selectbox("Eliminar nodo", [""] + graph["nodes"], key="remove_node")
                if node_to_remove and st.button("Eliminar Nodo"):
                    # Elimina el nodo y todas sus relaciones
                    graph["nodes"].remove(node_to_remove)
                    graph["edges"] = [e for e in graph["edges"] if e[0] != node_to_remove and e[1] != node_to_remove]

            st.markdown("#### Flechas (Relaciones)")
            if len(graph["nodes"]) >= 2:
                with st.form("add_edge_form"):
                    col1, col2 = st.columns(2)
                    with col1:
                        from_node = st.selectbox("Desde", graph["nodes"], key="from_node")
                    with col2:
                        to_node = st.selectbox("Hacia", graph["nodes"], key="to_node")
                    add_edge = st.form_submit_button("Agregar Flecha")
                    if add_edge and from_node != to_node:
                        if (from_node, to_node) not in graph["edges"]:
                            graph["edges"].append((from_node, to_node))
                        else:
                            st.warning("Esa flecha ya existe.")

            if len(graph["edges"]) > 0:
                edge_to_remove = st.selectbox(
                    "Eliminar flecha",
                    [""] + [f"{a} → {b}" for a, b in graph["edges"]],
                    key="remove_edge"
                )
                if edge_to_remove and st.button("Eliminar Flecha"):
                     edge_tuple = tuple(edge_to_remove.split(" → "))
                     if edge_tuple in graph["edges"]:
                        graph["edges"].remove(edge_tuple)

            # Actualiza el estado tras cambios
            st.session_state[user_key] = graph

         # Mostrar el grafo con graphviz
        dot = "digraph {\n"
        for node in graph["nodes"]:
            dot += f'    "{node}" [shape=ellipse];\n'
        for a, b in graph["edges"]:
            dot += f'    "{a}" -> "{b}";\n'
        dot += "}"
        st.graphviz_chart(dot)

    def _workload_monitor(self):
        with st.container(border=True):
            st.subheader("⚖️ Monitor de Carga de Trabajo")
            current_load = st.slider("Tu carga actual (1-10)", 1, 10, 7)
            ideal_load = st.slider("Carga ideal deseada (1-10)", 1, 10, 5)
            if current_load > ideal_load:
                st.error("¡Carga de trabajo excesiva!")
            elif current_load < ideal_load:
                st.warning("Capacidad disponible")
            else:
                st.success("Carga equilibrada")

    def _gamification_system(self):
        with st.container(border=True):
            st.subheader("🎮 Sistema de Recompensas")
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("🏅 Puntos Acumulados", "1,250")
            with col2:
                st.metric("🌟 Nivel Actual", "5")
            with col3:
                st.metric("🏆 Insignias", "3/10")

    def _learning_portal(self):
        st.subheader("🎓 Plataforma de Aprendizaje")

        user_key = f"learning_courses_{st.session_state.current_user}"
        if user_key not in st.session_state:
            st.session_state[user_key] = [
                {
                    "id": 1,
                    "title": "Gestión del Tiempo",
                    "progress": 0.4,
                    "url": "https://www.udemy.com/course/gestion-del-tiempo/"
                },
               {
                    "id": 2,
                    "title": "Liderazgo Efectivo",
                    "progress": 0.7,
                    "url": "https://www.linkedin.com/learning/search?entityType=COURSE&keywords=liderazgo&language=es_ES"
                },
               {
                    "id": 3,
                    "title": "Inteligencia Emocional",
                    "progress": 0.2,
                    "url": "https://www.linkedin.com/learning/search?entityType=COURSE&keywords=inteligencia&language=es_ES"
               }
            ]

        expander = st.expander("➕ Agregar nuevo curso vinculado", expanded=False)
        with expander:
            with st.form("add_course_form", clear_on_submit=True):
                new_title = st.text_input("Nombre del curso")
                new_url = st.text_input("Enlace al curso (Udemy, Coursera, Linkedin, etc.)")
                submitted = st.form_submit_button("Agregar curso")
                if submitted:
                    if new_title and new_url:
                        next_id = max([c["id"] for c in st.session_state[user_key]] + [0]) + 1
                        st.session_state[user_key].append({
                            "id": next_id,
                            "title": new_title,
                            "progress": 0.0,
                            "url": new_url
                        })
                        st.success("¡Curso agregado!")
                    else:
                        st.warning("Completa todos los campos.")

        for course in st.session_state[user_key]:
            with st.container():
                st.markdown(f"**{course['title']}**")
                c1, c2 = st.columns([5,1])
                with c1:
                    st.progress(course["progress"])
                with c2:
                    st.markdown(
                        f"<span style='font-size:18px;color:#007bff;font-weight:bold'>{int(course['progress']*100)}%</span>",
                        unsafe_allow_html=True
                    )
                st.markdown(
                    f"[Ir al curso]({course['url']})",
                    unsafe_allow_html=True
                )
                new_prog = st.slider(
                    "Actualizar progreso (%)",
                    0, 100, int(course["progress"]*100),
                    key=f"prog_{course['id']}")
                if new_prog != int(course["progress"]*100):
                    course["progress"] = new_prog / 100.0
                    st.experimental_rerun()
                if st.button(f"Eliminar {course['title']}", key=f"del_{course['id']}"):
                    st.session_state[user_key] = [
                        c for c in st.session_state[user_key] if c["id"] != course["id"]
                    ]
                    st.experimental_rerun()
    
    def _show_compliance(self):
        with st.expander("⚖️ Auditoría Normativa", expanded=True):
            uploaded_file = st.file_uploader("Subir Documento", type=["txt", "docx", "pdf"])
            
            if uploaded_file:
                text = ""
                try:
                    if uploaded_file.type == "text/plain":
                        text = uploaded_file.getvalue().decode()
                    elif uploaded_file.type == "application/pdf":
                        import PyPDF2
                        reader = PyPDF2.PdfReader(uploaded_file)
                        text = "\n".join([page.extract_text() for page in reader.pages])
                    else:
                        from docx import Document
                        doc = Document(uploaded_file)
                        text = "\n".join([para.text for para in doc.paragraphs])
                except Exception as e:
                    st.error(f"Error al leer el archivo: {str(e)}")
                    return
                
                audit_result = self._audit_document(text)
                st.write("**Resultados de Auditoría:**")
                st.json(audit_result)

    def _audit_document(self, text):
        doc = self.nlp(text)
        return {
            'GDPR': any(token.text.lower() in ('datos personales', 'consentimiento') for token in doc),
            'SOX': any(token.text.lower() in ('control interno', 'auditoría financiera') for token in doc),
            'ISO27001': any(token.text.lower() in ('seguridad de la información', 'riesgos') for token in doc)
        }

    def _show_payment(self):
        st.header("📈 Planes EnterpriseFlow")
        cols = st.columns(3)
        
        with cols[0]:
            st.subheader("Básico")
            st.markdown("""
                - 10 usuarios
                - Soporte básico
                - Reportes estándar
                **Precio: $99/mes**
            """)
            if st.button("Elegir Básico", key="basico"):
                self._handle_subscription('basico')

        with cols[1]:
            st.subheader("Premium")
            st.markdown("""
                - 50 usuarios
                - Soporte prioritario
                - Reportes avanzados
                **Precio: $299/mes**
            """)
            if st.button("Elegir Premium", key="premium"):
                self._handle_subscription('premium')

        with cols[2]:
            st.subheader("Enterprise")
            st.markdown("""
                - Usuarios ilimitados
                - Soporte 24/7
                - Personalización
                **Precio: $999/mes**
            """)
            if st.button("Contactar Ventas", key="enterprise"):
                st.info("contacto@enterpriseflow.com")

    def _handle_subscription(self, plan: str):
        try:
            if not st.session_state.current_user:
                raise ValueError("Debe iniciar sesión primero")
            
            subscription_data = self.payment.create_subscription(
                customer_email=st.session_state.current_user,
                price_key=plan
            )
            
            if subscription_data.get('client_secret'):
                st.session_state.subscription = subscription_data
                self._show_payment_confirmation()
            else:
                st.error("Error al crear la suscripción")
        
        except Exception as e:
            st.error(f"Error en suscripción: {str(e)}")

    def _show_payment_confirmation(self):
        with st.form("payment-form"):
            st.write("Complete los datos de pago")
            card_number = st.text_input("Número de tarjeta")
            expiry = st.text_input("MM/AA")
            cvc = st.text_input("CVC")
            
            if st.form_submit_button("Confirmar Pago"):
                try:
                    st.success("Pago procesado exitosamente!")
                    st.session_state.subscription = None
                except Exception as e:
                    st.error(f"Error en pago: {str(e)}")

# Crear las tablas (ejecuta una sola vez al instalar tu app)
conn = sqlite3.connect("enterprise_flow.db")
c = conn.cursor()
c.execute("""
CREATE TABLE IF NOT EXISTS user_rewards (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    user_email TEXT NOT NULL,
    puntos INTEGER DEFAULT 0,
    nivel INTEGER DEFAULT 1,
    insignias INTEGER DEFAULT 0,
    tareas_completadas INTEGER DEFAULT 0,
    dias_constancia INTEGER DEFAULT 0,
    updated_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    UNIQUE(user_email)
)
""")
c.execute("""
CREATE TABLE IF NOT EXISTS personal_goals (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    user_email TEXT NOT NULL,
    goal_text TEXT NOT NULL,
    created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
    completed BOOLEAN DEFAULT 0
)
""")
conn.commit()
conn.close()

conn = sqlite3.connect("enterprise_flow.db")
c = conn.cursor()
c.execute("""
CREATE TABLE IF NOT EXISTS medical_records (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    user_email TEXT NOT NULL,
    patologia TEXT,
    enfermedades TEXT,
    embarazo BOOLEAN DEFAULT 0,
    observaciones TEXT,
    FOREIGN KEY(user_email) REFERENCES employees(user_email)
)
""")
c.execute("""
CREATE TABLE IF NOT EXISTS leave_requests (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    user_email TEXT NOT NULL,
    tipo_permiso TEXT,
    fecha_inicio DATE,
    fecha_fin DATE,
    estado TEXT DEFAULT 'pendiente',
    motivo TEXT,
    observaciones TEXT,
    FOREIGN KEY(user_email) REFERENCES employees(user_email)
)
""")
conn.commit()
conn.close()

if __name__ == "__main__":
    EnterpriseFlowApp()
    
    
